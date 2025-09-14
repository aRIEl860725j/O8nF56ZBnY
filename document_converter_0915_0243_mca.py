# 代码生成时间: 2025-09-15 02:43:15
import dash
import dash_core_components as dcc
import dash_html_components as html
from dash.dependencies import Input, Output
import docx
import pandas as pd
from io import BytesIO
import base64
import numpy as np

# 定义 Dash 应用程序
app = dash.Dash(__name__)

# 定义应用程序布局
app.layout = html.Div([
    html.H1("Document Converter"),
    dcc.Upload(
        id='upload-data',
        children=html.Button('Upload Document'),
        multiple=True
    ),
    html.Div(id='output-data-upload'),
    dcc.Download(id='download-docx'),
    dcc.Interval(
        id='interval-component',
        interval=1*1000,  # in milliseconds
        n_intervals=0
    )
])

# 回调函数，当上传文件时触发
@app.callback(
    Output('output-data-upload', 'children'),
    [Input('upload-data', 'contents')]
)
def update_output(contents):
    if contents is None:
        return html.Div([])
    else:
        # 将上传的文件内容转换为 DataFrame
        content_type, content_string = contents.split(',')
        decoded = base64.b64decode(content_string)
        try:
            # 尝试读取上传的文件
            # 假设上传的文件是 Excel 文件
            excel_file = BytesIO(decoded)
            df = pd.read_excel(excel_file)
            # 如果 Excel 文件读取成功，返回 DataFrame
            return html.Div([
                html.H5("Uploaded Excel File"),
                html.Table(
                    [html.Tr([html.Th(col) for col in df.columns])] +
                    [html.Tr([html.Td(df.iloc[i][col]) for col in df.columns]) for i in range(df.shape[0])]
                )
            ])
        except Exception as e:
            # 如果读取文件时出现错误，返回错误信息
            return html.Div([html.H6(str(e))])

# 回调函数，当上传文件时触发，用于下载转换后的文档
@app.callback(
    Output('download-docx', 'data'),
    [Input('upload-data', 'contents')]
)
def generate_download(contents):
    if contents is None:
        return None
    else:
        # 将上传的文件内容转换为 DataFrame
        content_type, content_string = contents.split(',')
        decoded = base64.b64decode(content_string)
        try:
            # 尝试读取上传的文件
            # 假设上传的文件是 Excel 文件
            excel_file = BytesIO(decoded)
            df = pd.read_excel(excel_file)
            # 将 DataFrame 转换为 DOCX 文件
            doc = docx.Document()
            doc.add_heading('Converted Document', 0)
            for i, row in df.iterrows():
                p = doc.add_paragraph()
                p.add_run(str(row))
            # 将 DOCX 文件保存为 BytesIO 对象
            output = BytesIO()
            doc.save(output)
            # 将 BytesIO 对象转换为 base64 编码的字符串
            output.seek(0)
            encoded = base64.b64encode(output.getvalue()).decode()
            # 返回 base64 编码的 DOCX 文件
            return dcc.send_data_frame(df.to_excel, 'converted_document.xlsx', as_attachment=True,
                                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                     filename='converted_document.docx')
        except Exception as e:
            # 如果读取文件时出现错误，返回错误信息
            return None

if __name__ == '__main__':
    app.run_server(debug=True)